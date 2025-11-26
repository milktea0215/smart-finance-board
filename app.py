from flask import Flask, jsonify, request, send_from_directory
from flask_cors import CORS
import pandas as pd
import traceback
import os
import glob
# === OpenAI 客製化：讀取 API Key 並建立 client ===
from openai import OpenAI
from dotenv import load_dotenv  # ✅ 若有 .env 可自動讀取

load_dotenv()  # ✅ 會自動讀取 .env 檔（可放 OPENAI_API_KEY）

# ✅ 正確方式：從環境變數讀取，不要放整段 key 名稱
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # ← 環境變數名稱是 OPENAI_API_KEY
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

# ✅ 防呆：有設定就建立 client，否則提示錯誤
if OPENAI_API_KEY:
    client = OpenAI(api_key=OPENAI_API_KEY)
    print("✅ 已載入 OPENAI_API_KEY，模型：", OPENAI_MODEL)
else:
    client = None
    print("❌ 未偵測到 OPENAI_API_KEY！請確認環境變數或 .env 已設定。")

# ----------------------------------------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app = Flask(__name__, static_folder=BASE_DIR, static_url_path='')
CORS(app)

# ===== 行業比較用：高科技產業指標輸出資料夾 & 指標欄位 =====
INDICATOR_OUTPUT_DIR = os.path.join(BASE_DIR, "高科技產業指標輸出")

INDICATOR_COLUMNS = [
    '負債占資產比率',
    '長期資金占不動產、廠房及設備比率',
    '流動比率',
    '速動比率',
    '利息保障倍數',
    '應收帳款週轉率',
    '不動產、廠房及設備週轉率',
    '存貨週轉率',
    '總資產週轉率',
    '資產報酬率',
    '純益率',
    '權益報酬率',
    '每股盈餘',
    '現金流量比率',
    '現金再投資比率',
    '現金流量允當比率 (%)',
    '營運槓桿度',
    '財務槓桿度',
]


EXCEL_PATH = os.path.join(BASE_DIR, "財務指標輸出.xlsx")

# 啟動時就讀檔，讀不到會直接在 Console 顯示錯誤
try:
    df = pd.read_excel(EXCEL_PATH)

# 🆕 分出不同類型資料
    avg_df = df[df["公司"].astype(str).str.contains("平均值")].copy()
    company_df = df[~df["公司"].astype(str).str.contains("平均值")].copy()

# 🆕 分出各年度平均與全期平均
    avg_by_year_df = avg_df[avg_df["公司"].astype(str).str.match(r"^\d{4} 平均值$")].copy()
    avg_overall_df = avg_df[avg_df["公司"] == "平均值"].copy()

    print(f"✅ 已載入 Excel：{EXCEL_PATH}")
    print(f"　公司資料：{len(company_df)} 筆，年度平均：{len(avg_by_year_df)} 筆，全期平均：{len(avg_overall_df)} 筆")

except Exception as e:
    print("❌ 載入 Excel 失敗：", e)
    traceback.print_exc()
    df = pd.DataFrame()

# ===== 行業比較用共用工具函式 =====

def get_industry_file_path(industry_name):
    """
    將行業名稱轉成對應的指標輸出檔路徑
    例：industry_name = '半導體' -> 高科技產業指標輸出/半導體指標輸出.xlsx
    """
    filename = f"{industry_name}指標輸出.xlsx"
    path = os.path.join(INDICATOR_OUTPUT_DIR, filename)
    if not os.path.exists(path):
        raise FileNotFoundError(f"找不到檔案：{path}")
    return path


def load_industry_df(industry_name):
    """讀取某行業的指標輸出檔"""
    path = get_industry_file_path(industry_name)
    df_ind = pd.read_excel(path)
    return df_ind


def get_year_average_row(df_ind, year):
    """
    取得某一年的「年度平均值」列。
    依 1.py 的設計：公司欄 = '2015 平均值'、'2016 平均值'...
    年份欄 = 2015、2016 ...
    """
    mask = (df_ind['公司'] == f"{year} 平均值") & (df_ind['年份'] == year)
    rows = df_ind[mask]
    if rows.empty:
        return None
    return rows.iloc[0]


def get_all_yearly_average_rows(df_ind):
    """
    取得所有年度平均值的列（排除最底下那個「平均值」總平均）。
    """
    mask = df_ind['公司'].astype(str).str.endswith("平均值") & (df_ind['公司'] != "平均值")
    return df_ind[mask].copy()


@app.route("/")
def home():
    return send_from_directory(BASE_DIR, "index.html")

@app.route("/ping")
def ping():
    return jsonify({"ok": True})

@app.route("/get_indicator")
def get_indicator():
    import json
    from flask import Response
    import numpy as np

    global company_df, avg_by_year_df, avg_overall_df

    try:
        q = (request.args.get("company") or "").strip()
        if not q:
            return jsonify({"error": "請提供 company 參數"}), 400

        norm = q.replace("'", "").replace(" ", "")

        if df.empty:
            return jsonify({"error": "伺服器未成功載入 Excel"}), 500

        # 🔍 模糊搜尋公司名稱
        mask = df["公司"].astype(str).str.replace(" ", "", regex=False).str.contains(norm, na=False)
        rows = df[mask]

        if rows.empty:
            return jsonify({"error": f"找不到公司：'{q}'"}), 404

        # ------------------------------
        # 🔧 清理數據：NaN、Inf、Timestamp
        # ------------------------------
        clean_rows = rows.copy()

        # NaN、inf、-inf → None
        clean_rows = clean_rows.replace([np.nan, np.inf, -np.inf], None)

        # Timestamp → YYYY-MM-DD
        for col in clean_rows.columns:
            if pd.api.types.is_datetime64_any_dtype(clean_rows[col]):
                clean_rows[col] = clean_rows[col].dt.strftime("%Y-%m-%d")

        # ------------------------------
        # 🧮 同產業平均值（年度與全期）
        # ------------------------------
        # 年度平均值：轉成 dict 陣列
        avg_by_year = []
        if not avg_by_year_df.empty:
            avg_by_year = avg_by_year_df.replace({np.nan: None}).to_dict(orient="records")

        # 全期平均值：取第一筆
        avg_overall = {}
        if not avg_overall_df.empty:
            avg_overall = avg_overall_df.iloc[0].replace({np.nan: None}).to_dict()


        # ------------------------------
        # 🔰 組合完整 JSON
        # ------------------------------
        result = {
            "company": q,
            "data": clean_rows.to_dict(orient="records"),
            "average_by_year": avg_by_year,
            "average_overall": avg_overall
        }


        return jsonify(result)

    except Exception as e:
        import traceback
        print("❌ /get_indicator 錯誤：", e)
        traceback.print_exc()
        return jsonify({
            "error": "伺服器內部錯誤",
            "detail": str(e)
        }), 500

@app.route("/chatgpt_advice", methods=["POST"])
def chatgpt_advice():
    try:
        print("🧠 /chatgpt_advice 收到請求：", request.get_json(silent=True))
        if client is None:
            print("❌ client is None")
            return jsonify({"error": "後端未設定 OPENAI_API_KEY"}), 500

        payload = request.get_json(silent=True) or {}
        company = (payload.get("company") or "該公司").strip()
        indicators = payload.get("indicators") or {}

        # 僅取有數值的指標
        numeric_items = {k: v for k, v in indicators.items() if isinstance(v, (int, float))}

        prompt = (
            f"請根據以下指標數值，為公司「{company}」提供簡短中文財務分析建議：\n"
            + "\n".join([f"- {k}: {v}" for k, v in numeric_items.items()])
        )
        print("🧠 準備送出 prompt：", prompt[:200])

        completion = client.chat.completions.create(
            model=OPENAI_MODEL,
            temperature=0.4,
            messages=[
                {"role": "system", "content": "你是謹慎的財務分析專家，輸出簡潔中文建議。"},
                {"role": "user", "content": prompt},
            ],
        )
        advice = completion.choices[0].message.content.strip()
        print("✅ ChatGPT 回傳內容：", advice[:150])
        return jsonify({"advice": advice})

    except Exception as e:
        print("❌ /chatgpt_advice 錯誤：", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ====================== 🧠 多年度 AI 分析（正式報告版＋舉例說明） ======================
@app.route("/chatgpt_multi_advice", methods=["POST"])
def chatgpt_multi_advice():
    try:
        payload = request.get_json(silent=True) or {}
        company = (payload.get("company") or "該公司").strip()
        all_years_data = payload.get("all_years") or []

        if not all_years_data:
            return jsonify({"error": "缺少年度資料"}), 400
        if client is None:
            return jsonify({"error": "後端未設定 OPENAI_API_KEY"}), 500

        annual_advice = {}
        for row in all_years_data:
            year = str(row.get("年份") or str(row.get("年月"))[:4])
            indicators = {k: v for k, v in row.items() if isinstance(v, (int, float))}

            # 🧩 改版 prompt：以正式報告語氣撰寫 + 結尾附舉例
            prompt = f"""
你是一位專業的財務分析師，請根據以下 {year} 年公司「{company}」的財務指標，
撰寫一份中文財務分析建議報告，包含以下部分：

1. 開頭一句：「根據{year}年的財務指標，對{company}的財務分析建議如下：」
2. 條列式分析（請以「1.」「2.」「3.」等格式），每條說明一個面向（如資產效率、負債管理、現金流、獲利能力等），
   並在句中直接引用指標名稱與數值，例如「負債占資產比率為28.38%，顯示公司負債管理良好」。
3. 最後加上總結段，簡述整體財務體質與未來方向。
4. 最後新增一段「舉例：」說明，提供可實際執行的具體建議，例如改善資金運用或強化風險控管策略。
5. 請勿使用 **粗體** 或 Markdown 標記，維持純文字格式。

以下是指標資料：
{indicators}
"""

            completion = client.chat.completions.create(
                model=OPENAI_MODEL,
                temperature=0.4,
                messages=[
                    {"role": "system", "content": "你是嚴謹且具專業判斷的財務分析專家，請撰寫正式報告風格的中文財務分析，每項建議後附具體舉例。"},
                    {"role": "user", "content": prompt},
                ],
            )
            text = completion.choices[0].message.content.strip()
            annual_advice[year] = text

        # 🧾 五年整體分析（2020–2024） — 同樣採用報告式 + 舉例結尾
        recent_years = sorted(annual_advice.keys())[-5:]
        summary_prompt = f"""
你是一位資深財務顧問，請根據公司「{company}」最近五年的財務建議（{', '.join(recent_years)}），
撰寫一份「2020–2024 整體財務分析建議」，格式與年度分析相同，包含：

1. 開頭句：「根據2020–2024年的財務趨勢，對{company}的整體財務分析建議如下：」
2. 條列三至五項主要財務重點（如現金流、負債結構、營運效率、獲利能力等）。
3. 結論段落總結整體財務健康狀況與未來發展方向。
4. 結尾請加一段「舉例：」提出實際可執行的建議（例如調整資本結構、提升資金運用效率等）。
5. 請勿使用 **粗體** 或 Markdown 標記。

請根據你的專業，輸出完整的中文文字分析報告。
"""

        completion_sum = client.chat.completions.create(
            model=OPENAI_MODEL,
            temperature=0.4,
            messages=[
                {"role": "system", "content": "你是資深財務顧問，撰寫正式中文財務報告。"},
                {"role": "user", "content": summary_prompt},
            ],
        )
        summary_text = completion_sum.choices[0].message.content.strip()

        return jsonify({
            "annual": annual_advice,
            "summary_5yr": summary_text
        })

    except Exception as e:
        print("❌ /chatgpt_multi_advice 錯誤：", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ====================== 🧠 行業比較 AI 建議 ======================
@app.route("/chatgpt_industry_advice", methods=["POST"])
def chatgpt_industry_advice():
    """
    針對「兩個行業」與「2015–2024 趨勢」產生 AI 建議：
    - 整體比較結論
    - 六大財務面向趨勢比較
    - 行業一優勢與風險
    - 行業二優勢與風險
    - 未來發展與決策方向
    """
    try:
        payload = request.get_json(force=True) or {}
        industry1 = (payload.get("industry1") or "").strip()
        industry2 = (payload.get("industry2") or "").strip()
        years_data = payload.get("years_data") or []

        if client is None:
            return jsonify({"error": "後端未設定 OPENAI_API_KEY"}), 500

        if not industry1 or not industry2 or not years_data:
            return jsonify({"error": "缺少 industry1、industry2 或 years_data"}), 400

        # 為了避免 prompt 太長，可視情況只取最近 10 年
        # years_data 格式預期為：
        # [
        #   {
        #     "year": 2015,
        #     "industry1": {"負債占資產比率": 0.35, "ROE": 0.12, ...},
        #     "industry2": {...}
        #   },
        #   ...
        # ]
        # 這個 years_data 會由前端整理後送進來
        import json
        years_json_str = json.dumps(years_data, ensure_ascii=False)

        prompt = f"""
你是一位熟悉財務指標與產業分析的資深財務顧問。

現在有兩個行業：「{industry1}」與「{industry2}」，提供了 2015–2024（或部分年度）的主要財務指標資料。
資料格式如下（JSON 陣列，每一筆是一個年度）：

years_data = {years_json_str}

其中：
- year：西元年度，例如 2015、2016。
- industry1：該年度「{industry1}」行業的各項財務指標數值。
- industry2：該年度「{industry2}」行業的各項財務指標數值。
- 指標名稱可能包含：負債占資產比率、長期資金占不動產、流動比率、速動比率、利息保障倍數、
  應收帳款週轉率、存貨週轉率、總資產週轉率、資產報酬率、純益率、權益報酬率、每股盈餘、
  現金流量比率、現金再投資比率、現金流量允當比率、營運槓桿度、財務槓桿度… 等。

請根據這些數字的「水準」與「變化趨勢」，撰寫一份中文「行業比較 AI 建議報告」，輸出格式請嚴格依照下列 5 大段落：

一、整體比較結論
- 用 2～3 句話總結哪個行業整體財務體質較佳，以及主要原因。
- 可簡單說明：成長性、穩定度、風險高低。

二、六大財務面向趨勢比較
請依照下列六個面向，比較 10 年（或可用年度）中兩個行業的趨勢與差異：
1. 財務結構（如負債占資產比率、長期資金占不動產比率）
2. 償債能力（如流動比率、速動比率、利息保障倍數）
3. 營運效率／經營績效（如應收帳款週轉率、存貨週轉率、總資產週轉率）
4. 獲利能力（如資產報酬率 ROA、權益報酬率 ROE、純益率、每股盈餘）
5. 現金流量（現金流量比率、現金再投資比率、現金流量允當比率）
6. 槓桿度與風險（營運槓桿度、財務槓桿度）

每一個面向：
- 說明哪一個行業在大多數年度表現較好。
- 若近五年與前五年有明顯差異，可簡單描述「有改善」或「有惡化」。

三、「{industry1}」的優勢與風險
- 條列 3～5 點說明 {industry1} 行業在財務結構、獲利能力、現金流、槓桿風險等方面的優缺點。
- 若有波動較大的指標，請提醒可能的風險。

四、「{industry2}」的優勢與風險
- 條列 3～5 點說明 {industry2} 行業的優勢與風險，寫法與上一段相同。
- 避免與上一段用語完全重複。

五、未來 3～5 年發展與決策建議
- 綜合上述趨勢，說明兩個行業未來 3～5 年可能的發展方向（成長性、穩定性、風險）。
- 給出 2～4 個具體決策建議，例如：
  - 偏好成長型投資時較適合哪一個行業？
  - 偏好穩健／防禦型配置時應該注意哪一個行業的哪些風險？
  - 若是撰寫報告或研究，可以聚焦在哪幾個指標做比較。

注意事項：
- 全文請使用正式但淺顯易懂的中文。
- 不要輸出任何 Markdown 標記（例如 #、*、** 等），直接輸出純文字分段即可。
- 不要捏造不存在的數字，只需根據給定的趨勢方向與相對高低來評論即可。
"""

        completion = client.chat.completions.create(
            model=OPENAI_MODEL,
            temperature=0.4,
            messages=[
                {
                    "role": "system",
                    "content": "你是嚴謹且善於解釋的財務顧問，會根據數字趨勢撰寫中文行業分析建議。"
                },
                {
                    "role": "user",
                    "content": prompt
                },
            ],
        )

        text = completion.choices[0].message.content.strip()
        return jsonify({"text": text})

    except Exception as e:
        import traceback
        print("❌ /chatgpt_industry_advice 錯誤：", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route("/list_companies")
def list_companies():
    try:
        if df.empty:
            return jsonify({"companies": []})
        companies = sorted(df["公司"].dropna().astype(str).unique().tolist())
        return jsonify({"companies": companies})
    except Exception as e:
        print("❌ /list_companies 錯誤：", e)
        traceback.print_exc()
        return jsonify({"companies": []}), 500

# ====================== 📊 行業比較 API ======================

@app.route("/industry_list", methods=["GET"])
def industry_list():
    """
    掃描高科技產業指標輸出資料夾，列出所有可用行業名稱。
    例如：半導體指標輸出.xlsx -> 半導體
    """
    if not os.path.exists(INDICATOR_OUTPUT_DIR):
        return jsonify({"industries": []})

    pattern = os.path.join(INDICATOR_OUTPUT_DIR, "*指標輸出.xlsx")
    files = glob.glob(pattern)

    industries = []
    for f in files:
        base = os.path.basename(f)
        # 去掉「指標輸出.xlsx」
        name = base.replace("指標輸出.xlsx", "")
        industries.append(name)

    industries = sorted(set(industries))
    return jsonify({"industries": industries})


@app.route("/industry_compare", methods=["GET"])
def industry_compare():
    """
    取得某年份兩個行業的各項「年度平均值」指標，用來顯示上方表格。
    參數：
      industry1, industry2, year
    """
    industry1 = request.args.get("industry1")
    industry2 = request.args.get("industry2")
    year_str = request.args.get("year")

    if not industry1 or not industry2 or not year_str:
        return jsonify({"error": "缺少參數"}), 400

    try:
        year = int(year_str)
    except ValueError:
        return jsonify({"error": "年份格式錯誤"}), 400

    try:
        df1 = load_industry_df(industry1)
        df2 = load_industry_df(industry2)
    except FileNotFoundError as e:
        return jsonify({"error": str(e)}), 400

    row1 = get_year_average_row(df1, year)
    row2 = get_year_average_row(df2, year)

    if row1 is None or row2 is None:
        return jsonify({"error": f"{year} 年其中一個行業沒有年度平均值資料"}), 404

    indicators = []
    for col in INDICATOR_COLUMNS:
        val1 = row1.get(col, None)
        val2 = row2.get(col, None)

        # 轉成 float 或 None
        try:
            v1 = None if pd.isna(val1) else float(val1)
        except Exception:
            v1 = None
        try:
            v2 = None if pd.isna(val2) else float(val2)
        except Exception:
            v2 = None

        indicators.append({
            "name": col,
            "industry1": v1,
            "industry2": v2,
        })

    return jsonify({
        "year": year,
        "industry1": industry1,
        "industry2": industry2,
        "indicators": indicators,
    })


@app.route("/industry_trend", methods=["GET"])
def industry_trend():
    """
    取得兩個行業某一指標的歷年「年度平均值」趨勢，
    用來畫下方折線圖。
    參數：
      industry1, industry2, indicator
    """
    industry1 = request.args.get("industry1")
    industry2 = request.args.get("industry2")
    indicator = request.args.get("indicator")

    if not industry1 or not industry2 or not indicator:
        return jsonify({"error": "缺少參數"}), 400

    if indicator not in INDICATOR_COLUMNS:
        return jsonify({"error": f"未知指標: {indicator}"}), 400

    try:
        df1 = load_industry_df(industry1)
        df2 = load_industry_df(industry2)
    except FileNotFoundError as e:
        return jsonify({"error": str(e)}), 400

    avg1 = get_all_yearly_average_rows(df1)
    avg2 = get_all_yearly_average_rows(df2)

    # 年份交集，確保兩邊都有資料
    years1 = set(avg1['年份'].dropna().astype(int).tolist())
    years2 = set(avg2['年份'].dropna().astype(int).tolist())
    years = sorted(list(years1 & years2))

    years_labels = []
    values1 = []
    values2 = []

    for y in years:
        row1 = avg1[avg1['年份'] == y]
        row2 = avg2[avg2['年份'] == y]
        if row1.empty or row2.empty:
            continue

        v1 = row1.iloc[0].get(indicator, None)
        v2 = row2.iloc[0].get(indicator, None)

        try:
            v1 = None if pd.isna(v1) else float(v1)
        except Exception:
            v1 = None
        try:
            v2 = None if pd.isna(v2) else float(v2)
        except Exception:
            v2 = None

        years_labels.append(str(y))
        values1.append(v1)
        values2.append(v2)

    return jsonify({
        "indicator": indicator,
        "industry1": {
            "name": industry1,
            "years": years_labels,
            "values": values1,
        },
        "industry2": {
            "name": industry2,
            "years": years_labels,
            "values": values2,
        },
    })


# ====================== 📤 匯出報告：Word 專業版（含表格＋折線圖＋AI建議） ======================
from datetime import datetime
from io import BytesIO
import base64

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except Exception:
    DOCX_OK = False

def _b64_to_bytes(b64):
    """將 base64 圖片字串轉換為 bytes"""
    if "," in b64:
        b64 = b64.split(",", 1)[1]
    return base64.b64decode(b64)

@app.route("/export_report", methods=["POST"])
def export_report():
    """生成正式 Word 報告（含表格、圖表與 AI 建議）"""
    if not DOCX_OK:
        return jsonify({"error": "伺服器未安裝 python-docx"}), 500

    try:
        payload = request.get_json(force=True) or {}
        company = (payload.get("company") or "未命名公司").strip()
        selected_year = payload.get("selected_year", "")  # ✅ 新增接收選擇的年度
        indicators = payload.get("indicators_summary") or []  # [{name,value,avg,light}]
        charts = payload.get("charts") or []                  # [{title,img_base64}]
        ai_annual = payload.get("ai_annual") or {}            # {"2015":"...", ...}
        ai_summary_5yr = payload.get("ai_summary_5yr") or ""  # str

        # 🔧 準備輸出目錄
        reports_dir = os.path.join(BASE_DIR, "reports")
        os.makedirs(reports_dir, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_name = f"{company}_財務分析報告_{ts}.docx"
        docx_path = os.path.join(reports_dir, docx_name)

        # ====================== 📝 產生 Word 報告 ======================
        doc = Document()

        # 封面頁
        # ✅ 標題置中＋顯示年度
        main_title = f"公司財務分析報告（{selected_year} 年）" if selected_year else "公司財務分析報告"
        heading = doc.add_heading(main_title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title = doc.add_paragraph(f"公司名稱：{company}")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p = doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("（本報告由智慧財報分析儀表板自動生成）").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # 第一部分：財務指標摘要
        doc.add_heading("第一部分：財務指標摘要", level=1)
        if indicators:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "指標名稱"
            hdr[1].text = "公司數值"
            hdr[2].text = "同產業平均"
            hdr[3].text = "等級判斷"

            light_map = {"green": "🟢", "yellow": "🟡", "red": "🔴", "gray": "⚪"}
            for row in indicators:
                cells = table.add_row().cells
                cells[0].text = str(row.get("name", ""))
                cells[1].text = str(row.get("value", ""))
                cells[2].text = str(row.get("avg", ""))
                cells[3].text = light_map.get(row.get("light", "gray"), "⚪")
        else:
            doc.add_paragraph("（無可用財務指標資料）")

        doc.add_page_break()

        # 第二部分：歷年財務趨勢圖
        doc.add_heading("第二部分：歷年財務趨勢圖", level=1)
        if charts:
            for i, chart in enumerate(charts, 1):
                title = chart.get("title", f"圖 {i}")
                img_b64 = chart.get("img_base64")
                if not img_b64:
                    continue
                doc.add_paragraph(f"{i}. {title}")
                img_data = _b64_to_bytes(img_b64)
                img_stream = BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(6.5))
                doc.add_paragraph("")
        else:
            doc.add_paragraph("（未提供趨勢圖資料）")

        doc.add_page_break()

        # ✅ 第三部分：AI 財務分析建議（只顯示所選年度）
        ai_title = f"第三部分：AI 財務分析建議（{selected_year} 年）" if selected_year else "第三部分：AI 財務分析建議"
        doc.add_heading(ai_title, level=1)

        if ai_annual and selected_year and selected_year in ai_annual:
            doc.add_paragraph(ai_annual[selected_year])
        elif not selected_year and ai_annual:
            for year in sorted(ai_annual.keys()):
                doc.add_heading(f"{year} 年財務分析建議", level=2)
                lines = str(ai_annual[year]).split("\n")
                for line in lines:
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("（無對應年度 AI 建議資料）")
        doc.add_page_break()

        # ✅ 第四部分：AI 財務分析整體建議（2020–2024）
        doc.add_heading("第四部分：AI 財務分析整體建議（2020–2024）", level=1)
        summary_text = payload.get("ai_summary_5yr", "").strip()
        if summary_text:
            lines = summary_text.split("\n")
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("（目前無五年整體分析資料）")

        # 📦 儲存報告
        doc.save(docx_path)

        # ✅ 回傳給前端可下載的相對路徑
        return jsonify({
            "word": f"reports/{docx_name}",
            "pdf": None  # 若要加 PDF 轉換，可再擴充
        })

    except Exception as e:
        print("❌ 匯出報告錯誤：", e)
        traceback.print_exc()
        return jsonify({"error": f"匯出報告失敗：{e}"}), 500

@app.route("/export_industry_report", methods=["POST"])
def export_industry_report():
    """生成『行業比較』Word 報告（含指標表格＋趨勢圖＋AI 建議）"""
    if not DOCX_OK:
        return jsonify({"error": "伺服器未安裝 python-docx"}), 500

    try:
        payload = request.get_json(force=True) or {}

        # ===== 1. 取前端傳來的資料 =====
        industry1 = (payload.get("industry1") or "行業一").strip()
        industry2 = (payload.get("industry2") or "行業二").strip()
        year = payload.get("year")  # 可以是字串或數字，都接受

        indicators = payload.get("indicators") or []   # [{name, industry1_value, industry1_light, industry2_value, industry2_light}]
        charts = payload.get("charts") or []           # [{title, img_base64}]
        ai_text = (payload.get("ai_text") or "").strip()  # 行業 AI 建議全文

        # ===== 2. 準備輸出檔案路徑 =====
        reports_dir = os.path.join(BASE_DIR, "reports")
        os.makedirs(reports_dir, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        year_str = f"{year}" if year else "全部年份"

        docx_name = f"{industry1}_vs_{industry2}_行業比較報告_{year_str}_{ts}.docx"
        docx_path = os.path.join(reports_dir, docx_name)

        # ===== 3. 建立 Word 文件 =====
        doc = Document()

        # ---------- 封面 ----------
        title_text = f"行業比較財務分析報告（{year_str}）"
        heading = doc.add_heading(title_text, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub_p = doc.add_paragraph(f"行業一：{industry1}    行業二：{industry2}")
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_p = doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("（本報告由智慧財務分析與決策系統自動生成）").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # ---------- 第一部分：年度行業平均值指標比較 ----------
        doc.add_heading("第一部分：年度行業平均值指標比較", level=1)
        if year:
            doc.add_paragraph(f"本段整理 {year} 年「{industry1}」與「{industry2}」之行業平均財務指標。")

        if indicators:
            # 建立表格：指標名稱｜行業一｜行業一等級｜行業二｜行業二等級
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "指標名稱"
            hdr[1].text = industry1
            hdr[2].text = f"{industry1} 等級"
            hdr[3].text = industry2
            hdr[4].text = f"{industry2} 等級"

            light_map = {"green": "🟢", "yellow": "🟡", "red": "🔴", "gray": "⚪"}

            for row in indicators:
                name = str(row.get("name", ""))
                v1 = str(row.get("industry1_value", ""))
                v2 = str(row.get("industry2_value", ""))
                l1 = light_map.get(row.get("industry1_light", "gray"), "⚪")
                l2 = light_map.get(row.get("industry2_light", "gray"), "⚪")

                cells = table.add_row().cells
                cells[0].text = name
                cells[1].text = v1
                cells[2].text = l1
                cells[3].text = v2
                cells[4].text = l2
        else:
            doc.add_paragraph("（目前無行業指標比較資料）")

        doc.add_page_break()

        # ---------- 第二部分：歷年行業趨勢比較圖 ----------
        doc.add_heading("第二部分：歷年行業趨勢比較圖", level=1)
        if charts:
            for i, chart in enumerate(charts, start=1):
                title = chart.get("title", f"圖 {i}")
                img_b64 = chart.get("img_base64")
                if not img_b64:
                    continue

                doc.add_paragraph(f"{i}. {title}")
                img_data = _b64_to_bytes(img_b64)
                img_stream = BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(6.5))
                doc.add_paragraph("")  # 空一行
        else:
            doc.add_paragraph("（尚未提供行業趨勢圖資料）")

        doc.add_page_break()

        # ---------- 第三部分：AI 行業比較分析建議 ----------
        doc.add_heading("第三部分：AI 行業比較分析建議", level=1)

        if ai_text:
            # 保留原本換行
            for line in ai_text.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph("")  # 空行
        else:
            doc.add_paragraph("（尚未產生 AI 行業比較建議，可於系統中點選「產生 AI 建議」後再匯出報告。）")

        # ===== 4. 儲存並回傳路徑 =====
        doc.save(docx_path)

        return jsonify({
            "word": f"reports/{docx_name}"
        }), 200

    except Exception as e:
        print("❌ export_industry_report 錯誤：", e)
        traceback.print_exc()
        return jsonify({"error": f"匯出行業比較報告失敗：{e}"}), 500


if __name__ == "__main__":
    print("✅ Flask 啟動：http://127.0.0.1:5000")
    app.run(host="0.0.0.0", port=5000, debug=True)
